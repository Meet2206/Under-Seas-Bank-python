from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ABYSS = RGBColor(10, 18, 32)
DEPTH = RGBColor(22, 34, 60)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEAL = RGBColor(47, 224, 201)
GOLD = RGBColor(201, 162, 75)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E9FBF8"
PALE_GOLD = "FAF4E3"
TABLE_WIDTH_DXA = 9360


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc, title="Underseas Bank Documentation"):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.text = "Underseas Bank | Final Project Documentation"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=9, color=GRAY)


def add_title_block(doc, kicker, title, subtitle, metadata=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(kicker.upper())
    set_run_font(r, size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=26, color=ABYSS, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    set_run_font(r, size=13, color=GRAY)

    if metadata:
        for label, value in metadata:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            lr = p.add_run(f"{label}: ")
            set_run_font(lr, size=10.5, color=ABYSS, bold=True)
            vr = p.add_run(value)
            set_run_font(vr, size=10.5, color=GRAY)

    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="2FE0C9", size="10", space="6")
    rule.paragraph_format.space_after = Pt(12)


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_callout(doc, title, body, fill=PALE_TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=ABYSS, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=GRAY)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_GRAY):
    if widths is None:
        widths = [int(TABLE_WIDTH_DXA / len(headers))] * len(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=ABYSS, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if len(str(value)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=9.5, color=RGBColor(30, 41, 59))
    doc.add_paragraph()
    return table


def add_section_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_code_doc(path):
    doc = Document()
    configure_styles(doc, "Underseas Bank | Code and Imports")
    add_title_block(
        doc,
        "Technical Reference",
        "Code Explanation and Import Guide",
        "A module-by-module explanation of how the FastAPI backend, React frontend, security utilities, services, routes, and imports work together.",
        [("Status", "Updated final project documentation"), ("Scope", "Backend, frontend, services, utilities, routing, and design-system files")],
    )

    doc.add_heading("1. System Overview", level=1)
    add_para(doc, "Underseas Bank is a full-stack digital banking application. The backend is a FastAPI service responsible for authentication, account ownership, transactions, products, analytics, email/OTP behavior, persistence, and security. The frontend is a React/Vite application responsible for routing, user workflows, data display, dashboards, forms, and the final redesigned banking interface.")
    add_callout(doc, "Reading model", "Start with the runtime entry points, then models, schemas, routes, services, utilities, and finally the React pages. That order mirrors how data moves through the system.")

    doc.add_heading("2. Backend Entry and Configuration", level=1)
    add_table(doc, ["File", "Important imports", "Purpose"], [
        ("app/main.py", "FastAPI, CORSMiddleware, routers, Base/engine", "Creates the API app, configures CORS, registers routers, verifies database tables on startup, and exposes health endpoints."),
        ("app/config.py", "BaseSettings, lru_cache", "Centralizes environment-driven settings for database, JWT, Redis/Celery, OTP, email, encryption, app metadata, and CORS origins."),
        ("app/database.py", "create_engine, declarative_base, sessionmaker", "Defines the SQLAlchemy engine, base model class, and request-scoped database session provider."),
        ("app/redis_client.py", "redis, Optional", "Provides Redis connectivity used for OTP/cache/task-adjacent behavior when available."),
        ("app/celery_worker.py", "Celery, settings", "Configures background task execution through Redis broker/result backend."),
    ], [1900, 2600, 4860])

    doc.add_heading("3. Runtime Request Flow", level=1)
    add_numbered(doc, [
        "The browser renders a React page from `frontend/src/pages` through `App.jsx` routing.",
        "A user action calls a function in `frontend/src/services/api.js`.",
        "The API client attaches JSON headers and the bearer token for protected requests.",
        "FastAPI receives the request in a route module such as `transaction_routes.py`.",
        "Protected routes use `get_current_user` to decode the JWT and load the customer.",
        "The route validates ownership and delegates business logic to a service module.",
        "The service reads or writes SQLAlchemy models inside a database session.",
        "The route returns a schema-shaped JSON response, which React renders into the interface.",
    ])

    doc.add_heading("4. Data Models", level=1)
    add_table(doc, ["Model", "Stored data", "Why it matters"], [
        ("User", "Name, email, phone number, MPIN hash, verification and lock flags", "The ownership anchor for all customer accounts and products."),
        ("Account", "Encrypted account number, account type, user owner, balance", "Represents deposit accounts and the primary balance store."),
        ("Transaction", "Source/destination accounts, amount, type, timestamp/status", "Acts as the ledger and supports statements, passbooks, notifications, and analytics."),
        ("Loan", "Principal, rate, tenure, EMI, remaining balance, status", "Tracks borrowing products and repayment state."),
        ("FixedDeposit", "Principal, rate, duration, maturity amount, status", "Tracks term deposit investments and closure state."),
        ("CreditCard", "Encrypted card number, linked account, limit, used credit, status", "Tracks card issuance, usage, and bill payment."),
        ("Beneficiary", "Saved payee identity and encrypted account number", "Supports repeated transfers to known recipients."),
        ("Notification", "Customer activity metadata", "Provides a foundation for account activity alerts."),
    ], [1700, 3600, 4060])

    doc.add_heading("5. Routes and Services", level=1)
    add_table(doc, ["Route file", "Service layer", "Responsibility"], [
        ("auth_routes.py", "auth_service.py, otp_service.py", "Registration, login, current profile, OTP generation/verification, reset MPIN, and verification status."),
        ("account_routes.py", "account_service.py", "Create/list accounts, lookup account numbers, statements, and passbooks."),
        ("transaction_routes.py", "transaction_service.py", "Deposits, withdrawals, transfers, and transaction history."),
        ("loan_routes.py", "loan_service.py", "Loan application, active loans, and EMI payment."),
        ("fd_routes.py", "fd_service.py", "Fixed deposit creation, portfolio listing, and closure."),
        ("creditcard_routes.py", "creditcard_service.py", "Card application, card list, purchase simulation, and bill payment."),
        ("analytics_routes.py", "analytics modules", "Expense summary, category breakdown, financial health, and suggestions."),
        ("beneficiary_routes.py", "beneficiary model/service logic", "Saved recipient creation and listing."),
    ], [2300, 2600, 4460], PALE_BLUE)

    doc.add_heading("6. Security-Oriented Imports", level=1)
    add_table(doc, ["Utility", "Key dependency", "Explanation"], [
        ("password_hash.py", "Passlib bcrypt", "Hashes MPINs so the original PIN is not stored."),
        ("jwt_handler.py", "python-jose / JWT", "Creates and validates access tokens."),
        ("encryption.py", "cryptography.Fernet", "Encrypts account, card, and beneficiary identifiers at rest."),
        ("lookup_hash.py", "hashlib / HMAC-style lookup", "Allows exact-match searches without exposing encrypted source values."),
        ("otp_service.py", "secrets, Redis/fallback store", "Creates time-limited OTP codes and deletes them after successful verification."),
        ("auth_middleware.py", "OAuth2PasswordBearer, Depends", "Extracts bearer tokens and loads the current authenticated user."),
    ], [2100, 2300, 4960])

    doc.add_heading("7. Frontend Structure", level=1)
    add_table(doc, ["Area", "Files", "Role"], [
        ("Routing", "App.jsx", "Maps `/` to the landing page, `/login` to auth, and authenticated paths to banking pages."),
        ("API client", "services/api.js", "Centralizes fetch calls, attaches auth headers, parses errors, and exposes functions by domain."),
        ("Layout", "layout/MainLayout.jsx, Header.jsx, Sidebar.jsx", "Creates the app shell, mobile nav, notifications, profile menu, and sign-out flow."),
        ("Pages", "Landing, Login, Dashboard, Accounts, Transfer, Transactions, Loans, FixedDeposit, CreditCard, Analytics", "Implement user-facing workflows and data displays."),
        ("Components", "Card, ExpenseChart, TransactionTable", "Reusable UI and data-display components."),
        ("Design system", "index.css, auth.css, App.css", "Defines tokens, typography, responsive behavior, visual hierarchy, and the redesigned vault aesthetic."),
    ], [1800, 3400, 4160], PALE_BLUE)

    doc.add_heading("8. Design-System Files Added in the Final Pass", level=1)
    add_para(doc, "The final UI pass added a landing page and a tokenized redesign layer. The route `/` now presents the product landing page, while `/login` preserves the original authentication workflow. The CSS layer introduces abyss navy surfaces, bioluminescent teal accents, gold premium accents, focus rings, reduced-motion handling, vault cards, skeleton states, and polished marketing sections.")

    doc.add_heading("9. Configuration Reference", level=1)
    add_table(doc, ["Setting", "Purpose"], [
        ("DATABASE_URL", "PostgreSQL connection string used by SQLAlchemy."),
        ("SECRET_KEY / ALGORITHM / ACCESS_TOKEN_EXPIRE_MINUTES", "Controls JWT signing and access token lifetime."),
        ("REDIS_URL / CELERY_*", "Connects OTP/cache/background task infrastructure."),
        ("OTP_EXPIRE_SECONDS", "Defines OTP validity duration."),
        ("BREVO_API_KEY / SMTP fields", "Email delivery configuration. Brevo is the current primary API key path while SMTP fields remain as legacy reference."),
        ("FIELD_ENCRYPTION_KEY", "Fernet key for field-level encryption."),
        ("CORS_ORIGINS", "Allowed browser origins for frontend/API communication."),
        ("DEBUG", "Controls whether interactive API docs are exposed."),
    ], [2800, 6560])

    doc.add_heading("10. Current Strengths and Remaining Engineering Work", level=1)
    add_bullets(doc, [
        "The codebase separates models, schemas, routes, services, utilities, analytics, and frontend pages cleanly.",
        "Security primitives exist for MPIN hashing, JWT auth, OTP verification, field encryption, and ownership checks.",
        "The frontend now has a more cohesive visual hierarchy and a first-impression landing page.",
        "Future hardening should add Alembic migrations, stronger automated tests, decimal money handling, audit logging, refresh-token/session lifecycle improvements, and production secret rotation.",
    ])

    doc.save(path)


def build_logic_doc(path):
    doc = Document()
    configure_styles(doc, "Underseas Bank | Logic and Algorithms")
    add_title_block(
        doc,
        "Algorithm Reference",
        "Logic and Algorithms Used",
        "Plain-language explanation of the decision rules, calculations, security flows, and user workflows behind Underseas Bank.",
        [("Status", "Updated final project documentation"), ("Scope", "Authentication, OTP, authorization, transactions, products, analytics, and frontend state behavior")],
    )

    doc.add_heading("1. Authentication Logic", level=1)
    add_numbered(doc, [
        "Registration validates the submitted name, email, phone number, and four-digit MPIN through Pydantic schemas.",
        "The service checks uniqueness for identifiers that must not duplicate.",
        "The MPIN is hashed with bcrypt before persistence.",
        "The user record is saved and optional welcome/verification communication is triggered.",
        "Login validates credentials, checks lockout state, verifies the MPIN hash, and returns a signed JWT on success.",
        "Failed login attempts increment a counter and can lock the account after repeated failures.",
    ])

    doc.add_heading("2. OTP Algorithm", level=1)
    add_table(doc, ["Step", "Rule"], [
        ("Generate", "Create a six-digit code with a cryptographically safer random source."),
        ("Store", "Save the code with a short expiry window, preferring Redis where available."),
        ("Send", "Deliver the code through the configured email path."),
        ("Verify", "Compare the submitted value, reject expired/mismatched codes, and mark the target verified on success."),
        ("Delete", "Remove the code after success so it cannot be reused."),
    ], [1500, 7860])

    doc.add_heading("3. Authorization Algorithm", level=1)
    add_numbered(doc, [
        "Extract the bearer token from the Authorization header.",
        "Decode and validate the JWT signature and expiry.",
        "Load the user ID represented by the token payload.",
        "For account/product actions, load the target resource.",
        "Confirm the resource belongs to the current user before business logic runs.",
        "Reject missing, invalid, expired, or foreign-resource requests before state changes.",
    ])

    doc.add_heading("4. Money Movement Logic", level=1)
    add_table(doc, ["Operation", "Algorithm"], [
        ("Deposit", "Validate positive amount -> confirm account exists and belongs to user -> increase balance -> create deposit transaction -> commit."),
        ("Withdraw", "Validate positive amount -> confirm ownership -> verify sufficient balance -> decrease balance -> create withdrawal transaction -> commit."),
        ("Transfer", "Validate amount -> load source/destination -> confirm source ownership -> verify balance -> debit source -> credit destination -> create ledger transaction -> commit atomically."),
    ], [1800, 7560], PALE_BLUE)

    doc.add_heading("5. Transaction Consistency Rules", level=1)
    add_bullets(doc, [
        "Amounts must be positive before any balance change.",
        "Withdrawals and transfers must reject insufficient funds.",
        "Transfer debit and credit operations happen inside the same database unit of work.",
        "Failures should rollback the session so partial balance changes do not persist.",
        "Statements and passbooks derive from transaction history, not manually edited reports.",
    ])

    doc.add_heading("6. Account Statement and Passbook Logic", level=1)
    add_para(doc, "Statements and passbooks load transactions related to a chosen account and replay them in chronological order. Deposits increase the running balance; withdrawals decrease it; transfers increase or decrease the running balance depending on whether the account was sender or receiver.")

    doc.add_heading("7. Product Calculations", level=1)
    add_table(doc, ["Feature", "Formula / rule", "Output"], [
        ("Loan EMI", "P * r * (1 + r)^n / ((1 + r)^n - 1)", "Monthly installment amount."),
        ("Fixed deposit maturity", "principal * (1 + annual_rate * years)", "Estimated maturity amount."),
        ("Credit available", "credit_limit - used_credit", "Remaining purchasing power."),
        ("Financial health", "Savings behavior and spending signals", "Excellent / Good / Average / Poor style assessment."),
        ("Running balance", "Previous balance adjusted by transaction direction", "Statement/passbook balance after each event."),
    ], [2100, 4300, 2960])

    doc.add_heading("8. Loan Logic", level=1)
    add_bullets(doc, [
        "The selected disbursement account must belong to the current user.",
        "The requested principal, interest rate, and tenure are validated.",
        "The EMI calculator computes the monthly payment from principal, monthly interest, and tenure.",
        "A loan record stores original principal, EMI, remaining balance, and active/closed state.",
        "Payments reduce remaining balance and can move the loan toward closed status.",
    ])

    doc.add_heading("9. Fixed Deposit Logic", level=1)
    add_bullets(doc, [
        "The source account must belong to the current user.",
        "The deposit principal is deducted from the selected account if business rules require funding at creation.",
        "The maturity amount is calculated from principal, annual rate, and term length.",
        "Portfolio listing returns the user's own deposits only.",
        "Closure verifies ownership and updates deposit status.",
    ])

    doc.add_heading("10. Credit Card Logic", level=1)
    add_bullets(doc, [
        "A card application links the card to a user-owned account.",
        "Card numbers are treated as sensitive identifiers and encrypted at rest.",
        "Purchases increase used credit after checking available limit.",
        "Bill payment decreases used credit and restores available credit.",
        "Card status controls whether the product can be used.",
    ])

    doc.add_heading("11. Analytics Logic", level=1)
    add_para(doc, "Analytics collects transactions only for accounts owned by the authenticated customer. The system groups activity by category/type, computes spending summaries, produces chart-ready values for the frontend, and can generate rule-based financial tips or health scores.")

    doc.add_heading("12. Frontend State and Loading Logic", level=1)
    add_bullets(doc, [
        "Page components load data through `useEffect` and API helper functions.",
        "Dashboard data refreshes periodically to keep activity reasonably current.",
        "Protected API calls redirect unauthenticated users back to the login path.",
        "The final redesign adds skeleton display for loading account cards so the UI does not flash a false empty balance.",
        "The header collects recent transactions to populate a notification-style activity menu.",
    ])

    doc.add_heading("13. Error Handling and Failure Behavior", level=1)
    add_table(doc, ["Condition", "Expected response"], [
        ("Missing token", "Reject protected request and require login."),
        ("Expired/invalid token", "Clear token and require login again."),
        ("Zero or negative amount", "Reject request before balance mutation."),
        ("Insufficient balance", "Reject withdrawal or transfer."),
        ("Foreign account/product", "Reject access before returning sensitive data."),
        ("Invalid or expired OTP", "Reject verification and keep verification state unchanged."),
        ("Backend exception", "Return controlled 500 response with CORS headers instead of leaking traceback to browser."),
    ], [2600, 6760])

    doc.add_heading("14. End-to-End Scenario", level=1)
    add_numbered(doc, [
        "A visitor lands on the new marketing page and chooses Sign In or Open Account.",
        "The auth screen registers the customer or logs them in with MPIN.",
        "Email verification can require OTP confirmation before account access.",
        "The customer creates a savings/current/salary account.",
        "The customer deposits money, transfers funds, and reviews transaction history.",
        "The customer opens a loan, fixed deposit, or credit card product.",
        "Analytics summarizes activity and the dashboard displays account cards, insights, recent transactions, and spending visuals.",
    ])

    doc.add_heading("15. Future Algorithm Improvements", level=1)
    add_bullets(doc, [
        "Use Decimal for money calculations to avoid floating-point drift.",
        "Add idempotency keys for high-risk transfer endpoints.",
        "Add audit-log writes for authentication, product application, and money movement.",
        "Add scheduled jobs for FD maturity, loan due dates, card statement generation, and notification delivery.",
        "Add richer analytics based on merchant/category classification and historical trends.",
    ])

    doc.save(path)


def build_final_handbook(path):
    doc = Document()
    configure_styles(doc, "UNDERSEAS BANK | Final Handbook")
    add_title_block(
        doc,
        "Final Project Handbook",
        "UNDERSEAS BANK",
        "Complete explanation of the project from the smallest implementation details to the largest system architecture and user experience decisions.",
        [
            ("Prepared for", "Final project review and future handoff"),
            ("Stack", "FastAPI, PostgreSQL, SQLAlchemy, Redis, Celery, React 19, Vite, Recharts, Docker, Nginx"),
            ("Document type", "Comprehensive technical and product handbook"),
        ],
    )

    doc.add_heading("1. What Underseas Bank Is", level=1)
    add_para(doc, "Underseas Bank is a full-stack digital banking application that demonstrates secure customer authentication, account creation, deposits, withdrawals, transfers, statements, loans, fixed deposits, credit cards, beneficiaries, analytics, notifications, and a redesigned premium banking interface. It is structured as a realistic fintech-style project: a FastAPI backend owns business rules and persistence, while a React frontend owns user workflows and presentation.")
    add_callout(doc, "Core idea", "The app is not just a UI clone. It models the real responsibilities of a banking system: identity, authorization, ownership, ledger-style activity, product state, sensitive-data protection, and customer-facing clarity.", PALE_GOLD)

    doc.add_heading("2. Biggest-Picture Architecture", level=1)
    add_table(doc, ["Layer", "Technology", "Responsibility"], [
        ("Frontend", "React 19 + Vite + React Router + CSS", "Renders pages, captures form input, handles client routing, displays dashboards/charts, and calls the API."),
        ("API", "FastAPI", "Exposes route endpoints, validates authenticated access, and coordinates service logic."),
        ("Business logic", "Service modules", "Implements banking workflows such as transfers, deposits, loans, FDs, card purchases, and OTP verification."),
        ("Data validation", "Pydantic schemas", "Controls accepted request shape and returned response shape."),
        ("Persistence", "PostgreSQL + SQLAlchemy", "Stores users, accounts, transactions, products, beneficiaries, and notifications."),
        ("Async/cache", "Redis + Celery", "Supports background email/task behavior and OTP/cache-style workflows."),
        ("Security utilities", "bcrypt, JWT, Fernet, lookup hashes", "Protect MPINs, sessions, and sensitive identifiers."),
        ("Deployment support", "Docker, Docker Compose, Nginx, Render/Vercel config", "Supports local and hosted deployment paths."),
    ], [1600, 2600, 5160], PALE_BLUE)

    doc.add_heading("3. Repository Map", level=1)
    add_table(doc, ["Folder/file", "What it contains", "Why it exists"], [
        ("backend/app", "FastAPI application code", "The backend's executable core."),
        ("backend/app/models", "SQLAlchemy models", "Defines database tables and relationships."),
        ("backend/app/schemas", "Pydantic schemas", "Defines request/response contracts."),
        ("backend/app/routes", "FastAPI routers", "Defines public API endpoints."),
        ("backend/app/services", "Business logic", "Keeps route handlers thin and reusable."),
        ("backend/app/utils", "Security/math helpers", "Contains JWT, hashing, encryption, lookup hash, and EMI code."),
        ("backend/app/analytics", "Financial-intelligence modules", "Calculates expense analysis, health score, and tips."),
        ("frontend/src/pages", "React page components", "Implements each user-facing screen."),
        ("frontend/src/components", "Reusable UI pieces", "Contains header, sidebar, cards, charts, and transaction table."),
        ("frontend/src/services/api.js", "Central API client", "Keeps fetch logic and endpoint mapping in one place."),
        ("frontend/src/index.css", "Global design system", "Defines the final premium interface tokens and responsive behavior."),
        ("documents", "Generated project documentation", "Stores this final handbook and supporting docs."),
    ], [2200, 3200, 3960])

    doc.add_heading("4. Backend Runtime Entry", level=1)
    add_para(doc, "`backend/app/main.py` creates the FastAPI object, configures CORS, registers routers, exposes health endpoints, and verifies database tables during startup. The lifespan function retries database setup so startup is more tolerant when PostgreSQL is still booting. A global exception handler logs unexpected errors and returns a controlled JSON response with CORS headers so browser clients receive a real error instead of a misleading CORS failure.")

    doc.add_heading("5. Configuration and Environment", level=1)
    add_table(doc, ["Setting", "Meaning"], [
        ("DATABASE_URL", "Database location for PostgreSQL."),
        ("SECRET_KEY", "Private key used to sign JWTs. It must be replaced in production."),
        ("ACCESS_TOKEN_EXPIRE_MINUTES", "Controls how long access tokens remain valid."),
        ("REDIS_URL / CELERY_BROKER_URL / CELERY_RESULT_BACKEND", "Redis-backed cache/task infrastructure."),
        ("OTP_EXPIRE_SECONDS", "OTP expiry window, defaulting to five minutes."),
        ("BREVO_API_KEY", "Primary email API configuration path."),
        ("SMTP_* / EMAIL_FROM", "Legacy SMTP configuration retained for reference."),
        ("FIELD_ENCRYPTION_KEY", "Fernet key used to encrypt sensitive identifiers."),
        ("CORS_ORIGINS", "Comma/semicolon separated frontend origins allowed to call the API."),
        ("DEBUG", "Controls debug-only API documentation routes."),
    ], [3300, 6060], PALE_BLUE)

    doc.add_heading("6. Data Model Explained", level=1)
    add_para(doc, "The database is organized around ownership. A user owns accounts and products. Accounts generate transactions. Product tables such as loans, fixed deposits, and credit cards attach to a user and often a linked account. Beneficiaries attach to the user so saved recipients cannot leak across customers. Notifications represent activity messages or future alert records.")
    add_table(doc, ["Entity", "Tiny detail", "Bigger meaning"], [
        ("User", "Stores hashed MPIN and verification flags.", "Customer identity and access control anchor."),
        ("Account", "Stores encrypted account number and current balance.", "Primary money container."),
        ("Transaction", "Stores amount, type, source/destination and time.", "Ledger-style history for statements and analytics."),
        ("Loan", "Stores principal, EMI and remaining balance.", "Borrowing product lifecycle."),
        ("FixedDeposit", "Stores principal, duration and maturity amount.", "Investment product lifecycle."),
        ("CreditCard", "Stores encrypted card number, limit and used credit.", "Revolving credit lifecycle."),
        ("Beneficiary", "Stores saved recipient details.", "Convenience layer for repeated transfers."),
        ("Notification", "Stores activity alert data.", "Customer communication foundation."),
    ], [1800, 3400, 4160])

    doc.add_heading("7. Authentication in Detail", level=1)
    add_numbered(doc, [
        "The user submits registration or login data through the React auth page.",
        "The backend validates the payload through Pydantic schemas.",
        "Registration hashes the MPIN with bcrypt before saving.",
        "Login checks the submitted MPIN against the stored hash.",
        "Failed attempts can increase lockout counters.",
        "Successful login returns a signed JWT access token.",
        "The frontend stores the token and attaches it to future protected requests.",
        "Protected endpoints decode the token and load the current user before doing work.",
    ])

    doc.add_heading("8. OTP and Verification in Detail", level=1)
    add_para(doc, "OTP verification adds proof that the customer controls the email address or phone target. The implementation generates a six-digit code, stores it with an expiry, sends it through the configured communication path, verifies the submitted value, and removes the code on success. This prevents successful OTP reuse and keeps verification time-bound.")

    doc.add_heading("9. Authorization and Ownership", level=1)
    add_callout(doc, "Most important backend rule", "A token proves who the user is. It does not automatically prove that the requested account, loan, deposit, card, transaction history, or beneficiary belongs to that user. Each protected workflow must verify ownership before returning or changing data.")
    add_bullets(doc, [
        "Account lists are filtered by the current user.",
        "Transfers verify the source account belongs to the current user.",
        "Statements and passbooks verify the viewed account belongs to the current user.",
        "Loans, fixed deposits, cards, and beneficiaries are scoped to the current user.",
        "Analytics reads only transactions from the current user's accounts.",
    ])

    doc.add_heading("10. Money Movement", level=1)
    add_table(doc, ["Workflow", "Validation", "State change", "Record created"], [
        ("Deposit", "Amount positive; target account exists.", "Increase account balance.", "Deposit transaction."),
        ("Withdraw", "Amount positive; account exists; sufficient balance.", "Decrease account balance.", "Withdrawal transaction."),
        ("Transfer", "Amount positive; source owned; destination exists; sufficient balance.", "Debit source and credit destination.", "Transfer transaction."),
    ], [1700, 2860, 2600, 2200], PALE_BLUE)

    doc.add_heading("11. Statements and Passbooks", level=1)
    add_para(doc, "Statements and passbooks are derived views over transaction history. They are not separate manually maintained ledgers. The system loads all transactions related to an account and calculates how each event affects that account. This makes the generated history explainable and auditable.")

    doc.add_heading("12. Product Workflows", level=1)
    add_table(doc, ["Product", "What the user does", "Backend responsibility"], [
        ("Loan", "Choose account, amount, tenure and view estimated EMI.", "Validate ownership, compute EMI, store principal/rate/tenure/remaining balance."),
        ("Fixed Deposit", "Choose source account, amount, duration and view maturity value.", "Validate ownership, compute maturity amount, store term deposit lifecycle."),
        ("Credit Card", "Request a card and select a desired credit limit.", "Validate linked account, issue card record, track used and available credit."),
        ("Beneficiary", "Save recipient details.", "Encrypt sensitive recipient data and scope it to the customer."),
    ], [1700, 3400, 4260])

    doc.add_heading("13. Calculations and Formulas", level=1)
    add_table(doc, ["Calculation", "Formula", "Notes"], [
        ("Loan EMI", "P * r * (1 + r)^n / ((1 + r)^n - 1)", "P is principal, r is monthly interest, n is number of installments."),
        ("FD maturity", "principal * (1 + annual_rate * years)", "Simple-interest style maturity estimate."),
        ("Credit available", "credit_limit - used_credit", "Drives card spending availability."),
        ("Running balance", "previous balance +/- transaction effect", "Used by account history views."),
        ("Savings/health score", "Savings and spending signals", "Rule-based financial feedback."),
    ], [1900, 4200, 3260], PALE_GOLD)

    doc.add_heading("14. Analytics and Financial Intelligence", level=1)
    add_para(doc, "Analytics modules transform raw transactions into human-readable summaries. Expense analysis groups activity, financial health scores translate behavior into a simple rating, and the tips engine creates rule-based suggestions. On the frontend, Recharts turns the summary into chart components that users can scan quickly.")

    doc.add_heading("15. Frontend Routing", level=1)
    add_table(doc, ["Route", "Page", "Purpose"], [
        ("/", "Landing", "Marketing and first impression page introduced in the final redesign."),
        ("/login", "Login", "Authentication, registration, OTP, and reset MPIN flow."),
        ("/dashboard", "Dashboard", "Overview, account vault cards, shortcuts, insights, transactions, and spending chart."),
        ("/accounts", "Accounts", "Create accounts and review balances."),
        ("/transfer", "Transfer", "Deposit, withdraw, and transfer workflows."),
        ("/transactions", "Transactions", "Account transaction history."),
        ("/loans", "Loans", "Loan application and active loan review."),
        ("/fd", "FixedDeposit", "Fixed deposit creation and portfolio view."),
        ("/credit-card", "CreditCard", "Card application and existing card status."),
        ("/analytics", "Analytics", "Expense analysis and financial insights."),
    ], [1500, 2200, 5660], PALE_BLUE)

    doc.add_heading("16. Frontend Components", level=1)
    add_bullets(doc, [
        "`MainLayout` wraps authenticated pages with the sidebar and header.",
        "`Sidebar` provides navigation sections and sign-out.",
        "`Header` shows page title, activity notifications, profile menu, and mobile navigation trigger.",
        "`TransactionTable` renders recent account activity.",
        "`ExpenseChart` renders analytics visuals.",
        "`Card` provides a reusable visual container.",
        "`api.js` is the single source for frontend-to-backend endpoint calls.",
    ])

    doc.add_heading("17. Final UI Redesign", level=1)
    add_para(doc, "The final redesign applies the design-system direction called 'The Vault at Depth.' It uses deep navy backgrounds, disciplined teal bioluminescent accents, gold premium accents, Fraunces-style display typography, Geist/Inter-style UI typography, tokenized spacing, polished cards, skeleton loading states, accessible focus rings, and reduced-motion support. The design goal is serious premium banking, not playful ocean decoration.")
    add_table(doc, ["Design token idea", "Implementation meaning"], [
        ("Abyss navy", "Primary app shell, sidebar, vault cards, landing hero."),
        ("Bioluminescent teal", "Primary CTA, active navigation, focus rings, live/positive accents."),
        ("Gold", "Premium product cues such as credit card and fixed deposit surfaces."),
        ("Vault cards", "High-importance financial objects with deeper elevation and subtle motion."),
        ("Light data surfaces", "Forms and tables remain readable on white/light surfaces."),
        ("Reduced motion", "Animations are disabled or shortened for users who request less motion."),
    ], [2500, 6860])

    doc.add_heading("18. API Client Behavior", level=1)
    add_para(doc, "`frontend/src/services/api.js` defines the API base URL, wraps authenticated fetch calls, attaches bearer tokens, parses JSON, throws readable errors, and exposes domain-specific helpers for auth, accounts, transactions, loans, fixed deposits, credit cards, beneficiaries, analytics, and OTP. This keeps endpoint strings away from page components and makes frontend calls easier to maintain.")

    doc.add_heading("19. Security Details", level=1)
    add_table(doc, ["Risk", "Project response"], [
        ("Plaintext MPIN exposure", "Store bcrypt hash instead of original MPIN."),
        ("Token forgery", "Sign JWTs with SECRET_KEY and validate expiry."),
        ("Sensitive identifier exposure", "Encrypt account/card/beneficiary numbers at rest."),
        ("Lookup on encrypted data", "Use lookup hashes for exact-match queries."),
        ("OTP replay", "Expire codes and remove them on successful verification."),
        ("Cross-user data access", "Apply ownership checks before resource actions."),
        ("Browser CORS confusion", "Return controlled 500 responses with CORS headers."),
    ], [2600, 6760], PALE_BLUE)

    doc.add_heading("20. Deployment and Operations", level=1)
    add_bullets(doc, [
        "The backend can run locally with a Python virtual environment and Uvicorn.",
        "Docker Compose supports a multi-container backend stack with PostgreSQL, Redis, Celery, and Nginx.",
        "The frontend uses Vite and can run locally with `npm run dev` or build with `npm run build`.",
        "CORS origins must include the deployed frontend URL.",
        "Production deployments must replace default secrets and encryption keys.",
        "Database schema should eventually move to Alembic migrations rather than startup-time table creation.",
    ])

    doc.add_heading("21. Testing and Verification", level=1)
    add_para(doc, "The project includes a `backend/tests` package placeholder and frontend lint/build scripts. During the final UI pass, the frontend production build passed and lint passed with only one existing hook dependency warning in the transactions page. Future project hardening should add endpoint tests, authorization tests, transaction rollback tests, UI smoke tests, and regression tests for formulas.")

    doc.add_heading("22. Known Limitations", level=1)
    add_bullets(doc, [
        "Money values should eventually use Decimal rather than floating point.",
        "Database migrations are not yet formalized through Alembic in the current code layout.",
        "The project should add stronger test coverage before production use.",
        "Secrets in sample/default configuration must not be used as real production secrets.",
        "A complete banking platform would require compliance, audit, reconciliation, fraud detection, and operational monitoring beyond this showcase.",
    ])

    doc.add_heading("23. How to Explain the Project in an Interview", level=1)
    add_numbered(doc, [
        "Start by saying it is a full-stack digital banking showcase built with FastAPI and React.",
        "Explain the backend architecture: routers receive requests, schemas validate data, services perform business logic, models persist state.",
        "Explain the security model: MPIN hashing, JWT sessions, OTP verification, encrypted sensitive fields, and ownership checks.",
        "Explain the money workflows: deposit, withdraw, transfer, statement, passbook, loans, fixed deposits, and credit cards.",
        "Explain the frontend: routes, API client, authenticated layout, dashboard, charts, and redesigned premium UI.",
        "End with future improvements: Decimal money, tests, migrations, audit logs, refresh tokens, and production hardening.",
    ])

    doc.add_heading("24. Final Summary", level=1)
    add_para(doc, "Underseas Bank demonstrates a complete product story: a secure backend, a usable banking frontend, a modern design system, and documentation that explains the system from tiny code-level utilities to the full architecture. The most important engineering idea is separation of responsibility: routes expose the API, schemas validate data, services own business rules, models represent persistence, utilities handle security and math, and the frontend converts those capabilities into user workflows.")

    doc.save(path)


if __name__ == "__main__":
    build_code_doc("documents/Code_Explanation_and_Imports.docx")
    build_logic_doc("documents/Logic_and_Algorithms.docx")
    build_final_handbook("documents/UNDERSEAS BANK.docx")
